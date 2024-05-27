import os
import re
from typing import Optional, Any
from docx import Document
import hashlib
from embedchain.rag.nlp import tokenize_table, naive_merge, tokenize_chunks, rag_tokenizer

from langchain.text_splitter import RecursiveCharacterTextSplitter

from embedchain.chunkers.base_chunker import BaseChunker
from embedchain.config.add_config import ChunkerConfig
from embedchain.helpers.json_serializable import register_deserializable
from embedchain.deepdoc.parser.docx_parser import RAGFlowDocxParser


@register_deserializable
class DocxFileChunker(BaseChunker, RAGFlowDocxParser):
    """Chunker for .docx file."""

    def __init__(self, config: Optional[ChunkerConfig] = None):
        super().__init__(RAGFlowDocxParser)
        if config is None:
            config = ChunkerConfig(chunk_size=1000, chunk_overlap=0, length_function=len)
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.chunk_size,
            chunk_overlap=config.chunk_overlap,
            length_function=config.length_function,
        )
        super().__init__(text_splitter)

    @staticmethod
    def __clean(line):
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    @staticmethod
    def generate_doc_id(app_id, content):
        return str(app_id) + "-" + hashlib.sha256(content.encode()).hexdigest()

    def chunks(self, loader, src, metadata: Optional[dict[str, Any]] = None, config: Optional[ChunkerConfig] = None):
        self.doc = Document(src)
        min_chunk_size = config.min_chunk_size if config is not None else 1
        documents = []
        chunk_ids = []
        idMap = {}
        filename = os.path.basename(src)
        app_id = metadata.get("app_id", 1)
        knowledge_id = metadata.get("knowledge_id", 1)
        subject = metadata.get("subject", None)
        pn = 0
        from_page = int(metadata.get("from_page", 0))
        to_page = int(metadata.get("to_page", 100000))
        eng = metadata.get("eng", False)
        delimiter = metadata.get("delimiter", "\n!?。；！？")
        lines = []
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            if from_page <= pn < to_page and p.text.strip():
                lines.append(self.__clean(p.text))
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        tbls = []
        for tb in self.doc.tables:
            html= "<table>"
            for r in tb.rows:
                html += "<tr>"
                i = 0
                while i < len(r.cells):
                    span = 1
                    c = r.cells[i]
                    for j in range(i+1, len(r.cells)):
                        if c.text == r.cells[j].text:
                            span += 1
                            i = j
                    i += 1
                    html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                html += "</tr>"
            html += "</table>"
            tbls.append(((None, html), ""))
        sections = [(l, "") for l in lines if l]
        doc = {
            "docnm_kwd": filename,
            "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
        }
        doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
        cks = tokenize_table(tbls, doc, eng)
        chunks = naive_merge(sections, config.chunk_size, delimiter)
        cks.extend(tokenize_chunks(chunks, doc, eng, None))

        doc_id = self.generate_doc_id(app_id, "".join(each["content_with_weight"] for each in cks))
        metadatas = []
        for number, ck in enumerate(cks):
            chunk = ck["content_with_weight"]
            chunk_id = str(doc_id) + "-" + hashlib.sha256(chunk.encode()).hexdigest()
            meta_data = {}
            url = src
            # add data type to meta data to allow query using data type
            meta_data["app_id"] = app_id
            meta_data["doc_id"] = doc_id
            meta_data["knowledge_id"] = knowledge_id
            meta_data["hash"] = doc_id
            meta_data["data_type"] = self.data_type.value
            meta_data["subject"] = subject if subject is not None else os.path.basename(url)
            meta_data["status"] = 1
            meta_data['segment_number'] = number
            if idMap.get(chunk_id) is None and len(chunk) >= min_chunk_size:
                idMap[chunk_id] = True
                chunk_ids.append(chunk_id)
                documents.append(f"主题：{meta_data['subject']}。段落内容：{chunk}")
                metadatas.append(meta_data)
        return {
            "documents": documents,
            "ids": chunk_ids,
            "metadatas": metadatas,
            "doc_id": doc_id
        }
